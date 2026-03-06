"""Generate DATA_MAP.docx from structured data."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# --- Title ---
title = doc.add_heading('Data Map: Look-Alike & Behavioral Segmentation', level=0)
doc.add_paragraph(
    'Документ описывает все источники данных, необходимые для проекта, '
    'их статус и покрытие.\n'
    'Последнее обновление: 2026-03-06 (pilot run, Dec 2025, 10% sample)'
)

# --- Helper ---
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def status_color(status):
    if 'ПОЛУЧЕН' in status.upper():
        return RGBColor(0x22, 0x8B, 0x22)  # green
    elif 'НЕ ПОЛУЧЕН' in status.upper():
        return RGBColor(0xCC, 0x00, 0x00)  # red
    return RGBColor(0xFF, 0x8C, 0x00)  # orange

# --- Summary Table ---
doc.add_heading('Сводка по данным (pilot)', level=1)

summary_headers = ['Источник', 'Статус', 'Строк (pilot 10%)', 'Покрытие']
summary_rows = [
    ['Universe (paymentcounteragent_stran)', 'Получен', '1,443,852', '100%'],
    ['Базовые атрибуты (client_sdim)', 'Получен', '~1,443,852', '~100%'],
    ['Клиент/контрагент (clientinfo_shist)', 'Получен', '1,081,415', '74.9%'],
    ['ОКВЭД, город (redclients_shist_dmout)', 'Получен', '1,269,125', '87.9%'],
    ['Регион (client2smartsplace_shist)', 'Получен', '128,095', '8.9%'],
    ['Модельный сегмент (clientgeneralsegment)', 'Не получен', '0', '0%'],
    ['Выручка СПАРК (sparkfinindicator)', 'Не получен', '0', '0%'],
    ['Транзакции — поведение', 'В процессе', '—', '~100%'],
    ['Остатки (clientbalance_sagg)', 'В процессе', '—', '—'],
    ['Продукты (bmbpenetrproduct_stat)', 'В процессе', '—', '—'],
]
t = add_table(doc, summary_headers, summary_rows)
# Color status cells
for ri, row in enumerate(summary_rows):
    cell = t.rows[ri + 1].cells[1]
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = status_color(row[1])
            r.bold = True

doc.add_paragraph()

# --- 1. Universe ---
doc.add_heading('1. Universe — вселенная клиентов', level=1)
doc.add_paragraph('Источник: s_dmrb.paymentcounteragent_stran')
doc.add_paragraph('Поле: client_uk (DISTINCT)')
doc.add_paragraph(
    'Фильтры: date_part за декабрь 2025, deleted_flag != \'Y\', '
    'client_uk IS NOT NULL, hash-сэмпл 10%'
)
p = doc.add_paragraph('Статус: ')
r = p.add_run('ПОЛУЧЕН')
r.bold = True
r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
p.add_run(' — 1,443,852 уникальных client_uk')
doc.add_paragraph('Файл: data/universe.parquet')

# --- 2. Base ---
doc.add_heading('2. Базовые атрибуты клиента', level=1)
doc.add_paragraph('Источник: s_dmrb.client_sdim + clienttype_ldim')
doc.add_paragraph('JOIN: universe.client_uk = client_sdim.uk')
p = doc.add_paragraph('Статус: ')
r = p.add_run('ПОЛУЧЕН')
r.bold = True
r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)

add_table(doc,
    ['Поле', 'Описание', 'Получен'],
    [
        ['client_uk', 'Ключ клиента', 'Да'],
        ['client_name', 'Название/ФИО', 'Да'],
        ['entrepreneur_flag', 'Признак ИП', 'Да'],
        ['blacklist_flag', 'Чёрный список', 'Да'],
        ['client_type_name', 'Тип (ЮЛ/ФЛ/ИП)', 'Да'],
    ])

# --- 3. Client/Counterparty ---
doc.add_heading('3. Статус клиент / контрагент', level=1)
doc.add_paragraph('Источник: s_dmrb.clientinfo_shist (SCD-таблица)')
doc.add_paragraph('JOIN: clientinfo_shist.client_uk = universe.client_uk')
doc.add_paragraph('SCD-фильтр: effective_to >= \'2025-12-31\', ROW_NUMBER по effective_from DESC')
p = doc.add_paragraph('Статус: ')
r = p.add_run('ПОЛУЧЕН')
r.bold = True
r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
p.add_run(' — 1,081,415 строк (74.9% universe)')

doc.add_paragraph('Распределение:', style='List Bullet')
add_table(doc,
    ['clientcounterparty_flag', 'Кол-во', 'Описание'],
    [
        ['Y', '1,035,325', 'Клиенты банка'],
        ['N', '44,972', 'Контрагенты (проспекты)'],
        ['?', '1,118', 'Неизвестно'],
    ])
doc.add_paragraph()
add_table(doc,
    ['Поле', 'Описание', 'Получен'],
    [
        ['clientcounterparty_flag', 'Y/N/? — клиент или контрагент', 'Да'],
        ['clientchange_date', 'Дата перехода из контрагента в клиенты', 'Да'],
        ['counterpartychange_date', 'Дата перехода из клиента в контрагенты', 'Да'],
        ['client_rel_date', 'Дата начала отношений', 'Да'],
        ['registry_authcapital_amt', 'Уставный капитал (зарегистрированный)', 'Да'],
        ['paid_authcapital_amt', 'Уставный капитал (оплаченный)', 'Да'],
    ])

# --- 4. OKVED ---
doc.add_heading('4. ОКВЭД и город регистрации', level=1)
doc.add_paragraph('Источник: s_dmrb.dmoutdmrb_mdm_redclients_shist_dmout (SCD)')
doc.add_paragraph('JOIN: client_lp_uk = universe.client_uk (только ЮЛ)')
p = doc.add_paragraph('Статус: ')
r = p.add_run('ПОЛУЧЕН')
r.bold = True
r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
p.add_run(' — 1,269,125 строк (87.9% universe)')

add_table(doc,
    ['Поле', 'Описание', 'Получен'],
    [
        ['sparkokved_ccode', 'Код ОКВЭД из СПАРК', 'Да'],
        ['reg_city_name', 'Город регистрации', 'Да'],
        ['statusfu_ccode', 'Статус ФУ', 'Да'],
    ])
doc.add_paragraph('Примечание: Только для ЮЛ (client_lp_uk). ФЛ и ИП — NULL.')

# --- 5. Region ---
doc.add_heading('5. Регион (адрес)', level=1)
doc.add_paragraph('Источник: s_dmrb.client2smartsplace_shist (SCD)')
doc.add_paragraph('JOIN: client2smartsplace_shist.client_uk = universe.client_uk')
p = doc.add_paragraph('Статус: ')
r = p.add_run('ПОЛУЧЕН')
r.bold = True
r.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
p.add_run(' — 128,095 строк (8.9% universe)')

add_table(doc,
    ['Поле', 'Описание', 'Получен'],
    [
        ['addrref_region_uk', 'ID региона', 'Да'],
        ['addrref_city_name', 'Город', 'Да'],
    ])
p = doc.add_paragraph()
r = p.add_run('Низкое покрытие (8.9%)')
r.bold = True
p.add_run(' — таблица адресов заполнена не для всех. Город из п.4 (reg_city_name, 87.9%) как альтернатива.')

# --- 6. Segment ---
doc.add_heading('6. Модельный сегмент', level=1)
doc.add_paragraph('Источник: s_dmrb.dmrb_clientgeneralsegment_shist')
p = doc.add_paragraph('Статус: ')
r = p.add_run('НЕ ПОЛУЧЕН')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph(
    'Причина: client_uk в этой таблице хранится как DOUBLE (научная нотация, напр. 1.055E12) '
    'и не совпадает с universe.client_uk. Поле client_pin также не совпадает с client_sdim.client_pin. '
    'Таблица использует собственное пространство идентификаторов.'
)
doc.add_paragraph('Таблица содержит 141,724,857 строк.')

add_table(doc,
    ['Поле', 'Описание', 'Статус'],
    [
        ['srvpackagesegment_ccode', 'Модельный сегмент (Base, Alfa Smart, ...)', 'Не получен'],
        ['srvpackagelightsegment_ccode', 'Лёгкий сегмент', 'Не получен'],
        ['agesegment_ccode', 'Возрастной сегмент', 'Не получен'],
        ['client_active_flag', 'Флаг активности', 'Не получен'],
    ])
doc.add_paragraph(
    'Необходимо: таблица маппинга ID или документация DWH.',
    style='List Bullet'
)

# --- 7. Revenue ---
doc.add_heading('7. Выручка из СПАРК', level=1)
doc.add_paragraph('Источник: s_dmrb.sparkfinindicator_sstat через redclients_shist_dmout')
p = doc.add_paragraph('Статус: ')
r = p.add_run('НЕ ПОЛУЧЕН')
r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph(
    'Причина: Нет прямого FK. sparkcompany_src_ccode (STRING) → sparkcompany_uk (DOUBLE) '
    'через CAST не даёт совпадений. 0 rows matched.'
)

add_table(doc,
    ['Поле', 'Описание', 'Статус'],
    [
        ['revenue_rur_amt', 'Годовая выручка (RUR)', 'Не получен'],
        ['adjrevenue_rur_amt', 'Скорректированная выручка', 'Не получен'],
        ['failurescore_ccode', 'Скоринг надёжности', 'Не получен'],
        ['paymentindex_ccode', 'Платёжный индекс', 'Не получен'],
    ])
doc.add_paragraph(
    'Необходимо: правильный маппинг между redclients и sparkfinindicator.',
    style='List Bullet'
)

# --- 8. Transactions ---
doc.add_heading('8. Транзакционные агрегаты (поведение)', level=1)
doc.add_paragraph('Источник: s_dmrb.paymentcounteragent_stran')
doc.add_paragraph('Ноутбук: 02_etl_behavior.ipynb')
p = doc.add_paragraph('Статус: ')
r = p.add_run('В ПРОЦЕССЕ')
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

add_table(doc,
    ['Поле', 'Описание'],
    [
        ['tx_count', 'Количество транзакций'],
        ['total_amount', 'Суммарный оборот (RUR)'],
        ['total_outflow', 'Исходящие платежи'],
        ['total_inflow', 'Входящие платежи'],
        ['avg_tx_amount', 'Средняя сумма транзакции'],
        ['std_tx_amount', 'Стандартное отклонение суммы'],
        ['unique_counterparties', 'Уникальные контрагенты'],
        ['unique_payees', 'Уникальные получатели'],
        ['unique_payers', 'Уникальные плательщики'],
        ['active_months', 'Активные месяцы'],
        ['amount_first_half / second_half', 'Обороты 1-й/2-й половины периода'],
        ['cp_first_half / cp_second_half', 'Контрагенты 1-й/2-й половины'],
    ])

# --- 9. Balances ---
doc.add_heading('9. Остатки на счетах', level=1)
doc.add_paragraph('Источник: s_dmrb.clientbalance_sagg')
p = doc.add_paragraph('Статус: ')
r = p.add_run('В ПРОЦЕССЕ')
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

add_table(doc,
    ['Поле', 'Описание'],
    [
        ['avg_balance', 'Средний остаток за период'],
        ['max_balance', 'Максимальный остаток'],
        ['min_balance', 'Минимальный остаток'],
        ['avg_balance_30d', 'Средний 30-дневный остаток'],
    ])

# --- 10. Products ---
doc.add_heading('10. Продуктовое проникновение', level=1)
doc.add_paragraph('Источник: s_dmrb.bmbpenetrproduct_stat')
p = doc.add_paragraph('Статус: ')
r = p.add_run('В ПРОЦЕССЕ')
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

add_table(doc,
    ['Поле', 'Описание'],
    [
        ['product_count', 'Количество продуктов'],
        ['product_type_count', 'Количество типов продуктов'],
        ['product_total_amt', 'Суммарный объём по продуктам'],
    ])

# --- Derived features ---
doc.add_heading('Производные признаки (03_feature_engineering)', level=1)

add_table(doc,
    ['Признак', 'Формула', 'Зависит от'],
    [
        ['direction_ratio', 'total_outflow / total_amount', 'Транзакции'],
        ['avg_monthly_amount', 'total_amount / active_months', 'Транзакции'],
        ['amount_growth', '(2-я половина - 1-я) / 1-я', 'Транзакции'],
        ['cp_growth', 'Рост контрагентов аналогично', 'Транзакции'],
        ['tx_amount_cv', 'std / avg суммы транзакции', 'Транзакции'],
        ['okved_* (dummy)', 'One-hot top-20 ОКВЭД', 'ОКВЭД'],
        ['region_* (dummy)', 'One-hot top-20 регионов', 'Регион'],
        ['ctype_* (dummy)', 'One-hot тип клиента', 'Базовые атрибуты'],
    ])

# --- Blockers ---
doc.add_heading('Блокеры и рекомендации', level=1)

doc.add_heading('Критичные (для полной версии):', level=2)
doc.add_paragraph(
    'Модельный сегмент — нужна таблица маппинга ID '
    'dmrb_clientgeneralsegment_shist -> client_sdim (или документация DWH)',
    style='List Number'
)
doc.add_paragraph(
    'Выручка СПАРК — нужен правильный FK между '
    'redclients.sparkcompany_src_ccode и sparkfinindicator.sparkcompany_uk',
    style='List Number'
)

doc.add_heading('Рекомендации:', level=2)
doc.add_paragraph(
    'Регион — покрытие 8.9%. Использовать reg_city_name из ОКВЭД-таблицы (87.9%)',
    style='List Number'
)
doc.add_paragraph(
    'Ресурсы — для полного запуска нужны executor\'ы с >4GB heap '
    'или YARN-кластер с достаточными ресурсами',
    style='List Number'
)
doc.add_paragraph(
    'Период — pilot: 1 месяц + 10% sample. '
    'Для прода: SAMPLE_PCT=100, период 3-6 месяцев',
    style='List Number'
)

# --- Save ---
out_path = os.path.join(os.path.dirname(__file__), 'DATA_MAP.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
